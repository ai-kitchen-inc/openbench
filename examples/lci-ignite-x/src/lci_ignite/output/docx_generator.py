"""DOCX report generator for LCA analysis results.

Generates professionally formatted .docx reports from LCA analysis data
including IO tables, hotspot analysis, and narrative sections.
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openbench.core.abstractions import GeneratedOutput, OutputGenerator

logger = logging.getLogger(__name__)


class DocxReportGenerator(OutputGenerator):
    """Generate .docx reports from LCA analysis results.

    Creates structured reports with:
    - Title page
    - IO table section
    - Hotspot analysis section
    - Narrative explanations
    - PROPER 2025 references
    - Appendix

    Args:
        template: Optional template name (reserved for future use).
    """

    def __init__(self, template: str | None = None):
        self._template = template

    @property
    def output_format(self) -> str:
        return "docx"

    def validate(self, content: Any) -> bool:
        """Validate that content can be rendered as a .docx report."""
        if content is None:
            return False
        if isinstance(content, str):
            try:
                data = json.loads(content)
                return isinstance(data, dict)
            except (json.JSONDecodeError, TypeError):
                return bool(content.strip())
        return isinstance(content, dict)

    def generate(
        self,
        content: Any,
        template: str | None = None,
        output_path: str = "lca_report.docx",
        title: str | None = None,
        author: str | None = None,
        **options: Any,
    ) -> GeneratedOutput:
        """Generate a .docx report.

        Args:
            content: Report content. Can be:
                - dict with sections: {"io_table": ..., "hotspots": ..., "narrative": ...}
                - str (JSON or plain text)
            template: Template name (optional).
            output_path: Output file path.
            title: Report title.
            author: Report author.
            **options: Additional options.

        Returns:
            GeneratedOutput with file path and metadata.
        """
        # Parse content
        if isinstance(content, str):
            try:
                sections = json.loads(content)
            except (json.JSONDecodeError, TypeError):
                sections = {"narrative": content}
        elif isinstance(content, dict):
            sections = content
        else:
            sections = {"narrative": str(content)}

        # Create document
        doc = Document()

        # Title page
        report_title = title or sections.get("title", "LCA Analysis Report")
        self._add_title_page(doc, report_title, author)

        # IO Table section
        if "io_table" in sections:
            self._add_io_table_section(doc, sections["io_table"])

        # Hotspot analysis section
        if "hotspots" in sections:
            self._add_hotspot_section(doc, sections["hotspots"])

        # Narrative section
        if "narrative" in sections:
            self._add_narrative_section(doc, sections["narrative"])

        # PROPER 2025 references
        if "proper_references" in sections:
            self._add_proper_section(doc, sections["proper_references"])

        # Appendix
        if "appendix" in sections:
            self._add_appendix(doc, sections["appendix"])

        # Save
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        size = Path(output_path).stat().st_size

        return GeneratedOutput(
            file_path=output_path,
            format="docx",
            size_bytes=size,
            metadata={
                "title": report_title,
                "author": author,
                "sections": list(sections.keys()),
                "template": template or self._template,
            },
        )

    def _add_title_page(self, doc: Document, title: str, author: str | None) -> None:
        """Add title page."""
        doc.add_paragraph()  # Spacing
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if author:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Prepared by: {author}")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_page_break()

    def _add_io_table_section(self, doc: Document, io_data: Any) -> None:
        """Add IO table section."""
        doc.add_heading("Input-Output Tables", level=1)

        if isinstance(io_data, dict):
            for process_name, process_data in io_data.items():
                doc.add_heading(f"Process: {process_name}", level=2)

                inputs = process_data.get("inputs", [])
                outputs = process_data.get("outputs", [])

                if inputs:
                    doc.add_heading("Inputs", level=3)
                    table = doc.add_table(rows=1, cols=4)
                    table.style = "Table Grid"
                    headers = table.rows[0].cells
                    headers[0].text = "Flow"
                    headers[1].text = "Category"
                    headers[2].text = "Amount"
                    headers[3].text = "Unit"

                    for flow in inputs:
                        row = table.add_row().cells
                        row[0].text = str(flow.get("flow", ""))
                        row[1].text = str(flow.get("category", ""))
                        row[2].text = str(flow.get("amount", ""))
                        row[3].text = str(flow.get("unit", ""))

                if outputs:
                    doc.add_heading("Outputs", level=3)
                    table = doc.add_table(rows=1, cols=4)
                    table.style = "Table Grid"
                    headers = table.rows[0].cells
                    headers[0].text = "Flow"
                    headers[1].text = "Category"
                    headers[2].text = "Amount"
                    headers[3].text = "Unit"

                    for flow in outputs:
                        row = table.add_row().cells
                        row[0].text = str(flow.get("flow", ""))
                        row[1].text = str(flow.get("category", ""))
                        row[2].text = str(flow.get("amount", ""))
                        row[3].text = str(flow.get("unit", ""))

                doc.add_paragraph()
        elif isinstance(io_data, str):
            doc.add_paragraph(io_data)

    def _add_hotspot_section(self, doc: Document, hotspot_data: Any) -> None:
        """Add hotspot analysis section."""
        doc.add_heading("Environmental Hotspot Analysis", level=1)

        if isinstance(hotspot_data, dict):
            hotspots = hotspot_data.get("hotspots", [])
            if hotspots:
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                headers = table.rows[0].cells
                headers[0].text = "Rank"
                headers[1].text = "Flow"
                headers[2].text = "Amount"
                headers[3].text = "Impact %"
                headers[4].text = "Cumulative %"

                for i, h in enumerate(hotspots, 1):
                    row = table.add_row().cells
                    row[0].text = str(i)
                    row[1].text = str(h.get("name", ""))
                    row[2].text = f"{h.get('amount', 0)} {h.get('unit', '')}"
                    row[3].text = f"{h.get('percentage', 0):.1f}%"
                    row[4].text = f"{h.get('cumulative_percentage', 0):.1f}%"

            summary = hotspot_data.get("summary", "")
            if summary:
                doc.add_paragraph()
                doc.add_paragraph(summary)
        elif isinstance(hotspot_data, str):
            doc.add_paragraph(hotspot_data)

    def _add_narrative_section(self, doc: Document, narrative: Any) -> None:
        """Add narrative section."""
        doc.add_heading("Analysis Narrative", level=1)

        if isinstance(narrative, str):
            # Split by markdown headers
            for line in narrative.split("\n"):
                stripped = line.strip()
                if stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    doc.add_paragraph(stripped)
        elif isinstance(narrative, dict):
            for section_title, section_content in narrative.items():
                doc.add_heading(section_title, level=2)
                doc.add_paragraph(str(section_content))

    def _add_proper_section(self, doc: Document, references: Any) -> None:
        """Add PROPER 2025 references section."""
        doc.add_heading("PROPER 2025 References", level=1)

        if isinstance(references, list):
            for ref in references:
                if isinstance(ref, dict):
                    title = ref.get("title", "Reference")
                    content = ref.get("content", "")
                    doc.add_heading(title, level=2)
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph(str(ref))
        elif isinstance(references, str):
            doc.add_paragraph(references)

    def _add_appendix(self, doc: Document, appendix: Any) -> None:
        """Add appendix section."""
        doc.add_page_break()
        doc.add_heading("Appendix", level=1)

        if isinstance(appendix, str):
            doc.add_paragraph(appendix)
        elif isinstance(appendix, dict):
            for key, value in appendix.items():
                doc.add_heading(key, level=2)
                doc.add_paragraph(str(value))
