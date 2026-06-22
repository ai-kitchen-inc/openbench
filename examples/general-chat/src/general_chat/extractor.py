"""Document and image content extraction using Docling."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from openbench.chat.files import FileContentExtractor, StoredFile

logger = logging.getLogger(__name__)

_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
_DOCX_EXTENSIONS = {".docx", ".doc"}

_PDF_MIME_TYPES = {"application/pdf"}
_PDF_EXTENSIONS = {".pdf"}

_IMAGE_MIME_TYPES = {
    "image/png",
    "image/jpeg",
    "image/webp",
    "image/gif",
    "image/heic",
    "image/heif",
    "image/tiff",
    "image/bmp",
    "image/svg+xml",
}
_IMAGE_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
    ".heic",
    ".heif",
    ".tiff",
    ".tif",
    ".bmp",
    ".svg",
}

# PPTX/PPT have no lightweight extractor here, so they still use Docling.
_PPTX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "application/vnd.ms-powerpoint",
}
_PPTX_EXTENSIONS = {".pptx", ".ppt"}

# Heuristic: a digital PDF yields real text via pypdf; a scanned/image-only PDF
# yields ~nothing. Fall back to Docling OCR only when the text is essentially
# empty (tiny absolute floor + a small per-page floor to catch mostly-blank
# multi-page scans), so genuine text never pays the OCR cost.
_PDF_MIN_TOTAL_CHARS = 16
_PDF_MIN_CHARS_PER_PAGE = 4

_fallback = FileContentExtractor()

# Docling's DocumentConverter loads layout/table/OCR models on construction, so
# cache a single instance for the whole process instead of rebuilding per file.
_converter: Any | None = None


def _get_converter() -> Any:
    """Return a lazily-built, process-wide Docling DocumentConverter."""
    global _converter
    if _converter is None:
        from docling.document_converter import DocumentConverter

        _converter = DocumentConverter()
    return _converter


class DoclingContentExtractor:
    """Extract document text fast-first (pypdf/python-docx), with Docling as the
    OCR fallback for scanned PDFs and the primary path for PPTX/image OCR."""

    def extract(self, stored_file: StoredFile) -> str:
        ext = Path(stored_file.name).suffix.lower()
        # Fast path: digital PDFs via pypdf, only OCR via Docling when sparse.
        if _is_pdf(stored_file):
            return self._extract_pdf(stored_file)
        # Fast path: DOCX via python-docx, Docling only if it yields nothing.
        if stored_file.mime_type in _DOCX_MIME_TYPES or ext in _DOCX_EXTENSIONS:
            return self._extract_docx(stored_file)
        # PPTX/PPT: no lightweight extractor available — use Docling.
        if stored_file.mime_type in _PPTX_MIME_TYPES or ext in _PPTX_EXTENSIONS:
            return self._extract_with_docling(stored_file)
        return _fallback.extract(stored_file)

    def _extract_pdf(self, stored_file: StoredFile) -> str:
        """pypdf first; fall back to Docling OCR if the text layer is sparse."""
        text, pages = self._pdf_fast(stored_file)
        threshold = max(_PDF_MIN_TOTAL_CHARS, pages * _PDF_MIN_CHARS_PER_PAGE)
        if len(text.strip()) >= threshold:
            return text
        logger.info(
            "pypdf yielded sparse text for %s (%d chars, %d pages); trying Docling OCR",
            stored_file.name,
            len(text.strip()),
            pages,
        )
        ocr_text = self._extract_with_docling(stored_file, allow_pdf_fallback=False)
        if ocr_text.strip():
            return ocr_text
        if text.strip():
            return text
        return f"[{stored_file.name}] (document appears empty after extraction)"

    def _extract_docx(self, stored_file: StoredFile) -> str:
        text = self._extract_with_python_docx(stored_file)
        if text.strip() and "(document appears empty" not in text:
            return text
        # Empty result from python-docx — try Docling before giving up.
        docling_text = self._extract_with_docling(stored_file, allow_pdf_fallback=False)
        return docling_text if docling_text.strip() else text

    def extract_image(self, stored_file: StoredFile) -> dict[str, Any]:
        ext = Path(stored_file.name).suffix.lower()
        if stored_file.mime_type not in _IMAGE_MIME_TYPES and ext not in _IMAGE_EXTENSIONS:
            raise ValueError(f"Unsupported image source type: {ext or stored_file.mime_type}")

        from openbench.utils.media import is_svg, normalize_image, read_svg_text

        # Normalize HEIC/TIFF/BMP/GIF/SVG to a raster PNG/JPEG that Docling OCR
        # accepts. SVG markup is also captured directly as text.
        svg_text = read_svg_text(stored_file.path) if is_svg(stored_file.name, stored_file.mime_type) else ""
        try:
            ocr_path, _ = normalize_image(stored_file.path, stored_file.mime_type)
        except Exception as exc:
            logger.warning("Image normalization failed for %s: %s", stored_file.name, exc)
            ocr_path = stored_file.path

        try:
            result = _get_converter().convert(ocr_path)
            markdown = result.document.export_to_markdown().strip()
        except ImportError as exc:
            raise RuntimeError("docling is required for image OCR support.") from exc
        except Exception as exc:
            logger.warning("Docling image extraction failed for %s: %s", stored_file.name, exc)
            markdown = ""

        # SVG: the markup itself is meaningful text even when OCR finds nothing.
        if svg_text and not markdown:
            markdown = svg_text.strip()

        if not markdown:
            raise ValueError(f"No OCR text could be extracted from {stored_file.name}.")

        image_format = _image_format(stored_file.name, stored_file.mime_type)
        dimensions = _image_dimensions(ocr_path)
        width = dimensions.get("width")
        height = dimensions.get("height")
        description = _image_description(stored_file.name, image_format, width, height, markdown)
        search_text = _build_image_search_text(
            stored_file.name,
            image_format,
            description,
            markdown,
            dimensions,
        )

        return {
            "description": description,
            "ocr_text": markdown,
            "search_text": search_text,
            "metadata": {
                "format": image_format.lower(),
                "width": width,
                "height": height,
            },
        }

    def _extract_with_docling(
        self, stored_file: StoredFile, *, allow_pdf_fallback: bool = True
    ) -> str:
        try:
            result = _get_converter().convert(stored_file.path)
            text = result.document.export_to_markdown()
            if not text.strip():
                return f"[{stored_file.name}] (document appears empty after extraction)"
            return text
        except ImportError:
            # Docling not installed — degrade to the lightweight extractors.
            if allow_pdf_fallback and _is_pdf(stored_file):
                logger.warning("docling not installed; using pypdf for %s", stored_file.name)
                return self._pdf_fast(stored_file)[0]
            logger.warning("docling not installed; using python-docx for %s", stored_file.name)
            return self._extract_with_python_docx(stored_file)
        except Exception as exc:
            logger.warning("Docling extraction failed for %s: %s", stored_file.name, exc)
            ext = Path(stored_file.name).suffix.lower()
            if allow_pdf_fallback and _is_pdf(stored_file):
                return self._pdf_fast(stored_file)[0]
            if stored_file.mime_type in _DOCX_MIME_TYPES or ext in _DOCX_EXTENSIONS:
                return self._extract_with_python_docx(stored_file)
            return f"[{stored_file.name}] (extraction failed: {exc})"

    def _pdf_fast(self, stored_file: StoredFile) -> tuple[str, int]:
        """Extract PDF text with pypdf. Returns (text, page_count)."""
        try:
            from pypdf import PdfReader
        except ImportError:
            logger.warning("pypdf not installed; install it with: pip install pypdf")
            return "", 0

        try:
            reader = PdfReader(stored_file.path)
            pages = len(reader.pages)
            parts: list[str] = []
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    parts.append(f"### Page {index}\n\n{text}")
            return "\n\n".join(parts).strip(), pages
        except Exception as exc:
            logger.warning("pypdf extraction failed for %s: %s", stored_file.name, exc)
            return "", 0

    def _extract_with_python_docx(self, stored_file: StoredFile) -> str:
        try:
            import docx

            doc = docx.Document(stored_file.path)
            parts: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))

            full_text = "\n\n".join(parts)
            if not full_text.strip():
                return f"[{stored_file.name}] (document appears empty after extraction)"
            return full_text
        except ImportError:
            logger.warning("python-docx not installed; install it with: pip install python-docx")
            return _fallback.extract(stored_file)
        except Exception as exc:
            logger.warning(
                "python-docx extraction failed for %s: %s", stored_file.name, exc
            )
            return _fallback.extract(stored_file)


def _is_pdf(stored_file: StoredFile) -> bool:
    ext = Path(stored_file.name).suffix.lower()
    return stored_file.mime_type in _PDF_MIME_TYPES or ext in _PDF_EXTENSIONS


def _image_dimensions(path: str) -> dict[str, int | None]:
    try:
        from PIL import Image
    except ImportError:
        return {"width": None, "height": None}

    try:
        with Image.open(path) as image:
            width, height = image.size
    except Exception:
        return {"width": None, "height": None}

    return {"width": int(width), "height": int(height)}


def _image_format(name: str, mime_type: str) -> str:
    ext = Path(name).suffix.lower()
    if ext in {".jpg", ".jpeg"} or mime_type == "image/jpeg":
        return "JPEG"
    if ext == ".webp" or mime_type == "image/webp":
        return "WEBP"
    return "PNG"


def _image_description(
    name: str, image_format: str, width: int | None, height: int | None, ocr_text: str
) -> str:
    size = f"{width}x{height}" if width and height else "unknown size"
    if ocr_text.strip():
        return f"{image_format} image source {name} ({size}) with OCR-detected text."
    return f"{image_format} image source {name} ({size}) with no OCR text detected."


def _build_image_search_text(
    name: str,
    image_format: str,
    description: str,
    ocr_text: str,
    dimensions: dict[str, int | None],
) -> str:
    width = dimensions.get("width")
    height = dimensions.get("height")
    lines = [
        f"## {name}",
        "",
        "### Image summary",
        description,
        f"Format: {image_format}",
        f"Dimensions: {width or 'unknown'} x {height or 'unknown'}",
        "",
        "### Detected text",
        ocr_text.strip() or "(No OCR text detected.)",
    ]
    return "\n".join(lines).strip()
